#todo:   testing, fool-proofing
"""Copyright Marat Samigullin
This app creates a MS Word template of a service locating report according to parameters.
The name of the created file consists of the address and suburb name
The app also creates a template for the email to be sent to the client (file MS Word email template for the client) 
"""


from tkinter import *
from tkinter import filedialog, messagebox
from tkinter.ttk import Checkbutton, Combobox
import json	
from docx import Document
from docx.shared import Pt, Mm
from docx.shared import RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date
from cv2 import imread #pip install opencv-python



class Locating_report:

    def __init__(self):
        self.client_name=""
        self.images_paths=list()
              
    def start_window(self):
        self.main_window=Tk()
        # add widgets here
        self.main_window.title('Locating report generator')
        self.main_window.geometry("600x500")

        Intro_label = Label(self.main_window, text="Locating report generator - please fill in all the fields\n")
        Intro_label.pack( side=TOP, anchor=NW, padx=10)

        self.list_of_operators=self.templates["Operator_cert"]
        self.operator_combo = self.standard_combo("Operator",self.list_of_operators)
        
        

        today = date.today()
        self.date_entry = self.standard_entry("Date:",today.strftime("%d/%m/%Y"))
        self.client_name_entry = self.standard_entry("Client name:","Name of the client")       
        self.docket_number_entry = self.standard_entry("Docket#:","-")

        
        self.is_GPR=IntVar() #CheckButton works bugless only if the variable is of its own tkinter type (IntVar or others)
        self.is_GPR_tick = Checkbutton(self.main_window, text='GPR was used successfully on site', variable = self.is_GPR)
        self.is_GPR_tick.pack(side=TOP, anchor=NW, padx=10)
        
        self.is_GPR_bad_cond=IntVar()
        self.is_GPR_bad_cond_tick = Checkbutton(self.main_window, text='Bad soil conditions for GPR', variable = self.is_GPR_bad_cond)
        self.is_GPR_bad_cond_tick.pack(side=TOP, anchor=NW, padx=10)
        

        self.is_no_non_conductive=IntVar()
        self.is_no_non_conductive_tick = Checkbutton(self.main_window, text='No non-conductive was found', variable = self.is_no_non_conductive)
        self.is_no_non_conductive_tick.pack(side=TOP, anchor=NW, padx=10)


        self.is_pushrod=IntVar()
        self.is_pushrod_tick = Checkbutton(self.main_window, text='Push-rod was used on site', variable = self.is_pushrod)
        self.is_pushrod_tick.pack(side=TOP, anchor=NW, padx=10)

        

        self.suburb_name_entry=self.standard_entry("Suburb name:","Suburb name")

        self.address_entry=self.standard_entry("Address:","Address without a suburb")
       
       
        generate_button = Button(self.main_window,command = self.open_images, text="Choose the images for the report")
        generate_button.pack(side=TOP, anchor=NW, padx=10, pady=10)
        self.images_quantity = 0
        self.images_label = Label(self.main_window, text= f"{self.images_quantity} photos selected")
        self.images_label.pack( side=TOP, anchor=NW, padx=10)

        generate_button = Button(self.main_window,command = self.generate_report, text="Create the report")
        generate_button.pack(side=BOTTOM, anchor=SE, padx=10, pady=10)


        


        self.main_window.mainloop()

    def open_images(self):
        
        self.images_paths = filedialog.askopenfilenames(filetypes=[("Image files", ".jpeg .jpg .png .bmp .tiff .tif")])
        self.images_quantity = len(self.images_paths)
        self.images_label.config(text= f"{self.images_quantity} photos selected", fg='#008006')#green
        
        
        

    def standard_entry(self,label_text, default_text):
        my_frame=Frame(self.main_window)
        my_label = Label(my_frame, text=label_text)
        my_label.pack( side=LEFT, anchor=NW,padx=10)
        my_entry = Entry(my_frame, bd = 1)
        my_entry.insert(0,default_text)
        my_entry.pack(side=LEFT, anchor=NW, fill=X,  padx=10, expand=True)
        my_frame.pack(side=TOP,anchor=NW, fill=X, expand=False,pady=2)
        return my_entry

    def standard_combo(self,label_text, default_text):
            my_frame=Frame(self.main_window)
            my_label = Label(my_frame, text=label_text)
            my_label.pack( side=LEFT, anchor=NW,padx=10)
            my_combobox = Combobox(my_frame, values=default_text )
            
            my_combobox.pack(side=LEFT, anchor=NW, fill=X,  padx=10, expand=True)
            my_frame.pack(side=TOP,anchor=NW, fill=X, expand=False)
            return my_combobox


    def load_templates(self):
        filename = open("templates.json")
        self.templates=json.load(filename)
        return 0





    def set_cell_margins(self, table, left=0, right=0):

        tc = table._element
        tblPr = tc.tblPr
        tblCellMar = OxmlElement('w:tblCellMar')
        kwargs = {"left":left, "right":right}
        for m in ["left","right"]:
            node = OxmlElement("w:{}".format(m))
            node.set(qn('w:w'), str(kwargs.get(m)))
            node.set(qn('w:type'), 'dxa')
            tblCellMar.append(node)
        tblPr.append(tblCellMar)




    def generate_report(self): #invokes by pressing Generate Report button

        self.client_name=self.client_name_entry.get()
        self.suburb_name=self.suburb_name_entry.get()
        self.address=self.address_entry.get()
        self.date=self.date_entry.get()
        self.docket_number=self.docket_number_entry.get()


        try:
            self.operator_name = self.operator_combo.get().split("'")[1] #parses the operator's name
            self.operator_certificate_number = self.operator_combo.get().split("'")[3] #parses the certificate number
        except IndexError:
            self.operator_name = "experienced utility locator"
            self.operator_certificate_number = "121304"

        self.my_locating_report = Document()
        
        current_section = self.my_locating_report.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        current_section.orientation = WD_ORIENT.LANDSCAPE
        current_section.page_width = new_width
        current_section.page_height = new_height

        current_section.left_margin = Mm(10)
        current_section.right_margin = Mm(10)
        current_section.top_margin = Mm(1)
        current_section.bottom_margin = Mm(1)
        current_section.header_distance =Mm(10)
        current_section.footer_distance = Mm(10)
        current_section.gutter = Mm(0)
        



        #setting up a style
        obj_styles = self.my_locating_report.styles
        obj_parstyle = obj_styles.add_style('locating_report_style', WD_STYLE_TYPE.PARAGRAPH)
        obj_font = obj_parstyle.font
        obj_font.size = Pt(14)
        obj_font.name = 'Calibri (Body)'
        par_format= obj_parstyle.paragraph_format
        par_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        
        


         #setting up a style
        obj_styles = self.my_locating_report.styles
        obj_parstyle = obj_styles.add_style('locating_report_style_2', WD_STYLE_TYPE.PARAGRAPH)
        obj_font = obj_parstyle.font
        obj_font.size = Pt(11)
        obj_font.name = 'Calibri (Body)'
        par_format= obj_parstyle.paragraph_format.line_spacing = 1.15
        
        

        #setting up a title style
        obj_styles = self.my_locating_report.styles
        obj_parstyle = obj_styles.add_style('locating_report_title_style', WD_STYLE_TYPE.PARAGRAPH)
        obj_font = obj_parstyle.font
        obj_font.size = Pt(36)
        obj_font.name = 'Calibri (Body)'

        #setting up a header/footer style
        obj_styles = self.my_locating_report.styles
        obj_parstyle = obj_styles.add_style('locating_report_header_style', WD_STYLE_TYPE.TABLE)
        obj_parstyle.paragraph_format.space_after = Pt(2)
        obj_font = obj_parstyle.font
        obj_font.size = Pt(12)
        obj_font.name = 'Calibri (Body)'

        #setting up a header/footer style #2
        obj_styles = self.my_locating_report.styles
        obj_parstyle = obj_styles.add_style('locating_report_header_style_2', WD_STYLE_TYPE.PARAGRAPH)
        obj_parstyle.paragraph_format.space_after = Pt(1)
        obj_parstyle.paragraph_format.space_before = Pt(1)
        obj_font = obj_parstyle.font
        obj_font.size = Pt(8)
        obj_font.name = 'Calibri (Body)'


        #setting up a legend style
        obj_styles = self.my_locating_report.styles
        obj_parstyle = obj_styles.add_style('legend_style', WD_STYLE_TYPE.PARAGRAPH)
        obj_parstyle.paragraph_format.space_before = Pt(10)
        obj_parstyle.paragraph_format.space_after = Pt(0)
        obj_parstyle.paragraph_format.line_spacing = 1
        obj_font = obj_parstyle.font
        obj_font.size = Pt(9)
        obj_font.name = 'Calibri (Body)'


        #setting up a legend style_2 for representation of a coloured line
        obj_styles = self.my_locating_report.styles
        obj_parstyle = obj_styles.add_style('legend_style_2', WD_STYLE_TYPE.PARAGRAPH)
        obj_parstyle.paragraph_format.space_after = Pt(0)
        obj_parstyle.paragraph_format.space_before = Pt(0)
        obj_parstyle.paragraph_format.line_spacing = 1
        obj_parstyle.paragraph_format.left_indent = 0
        obj_parstyle.paragraph_format.first_line_indent = 0
        obj_font = obj_parstyle.font
        obj_font.size = Pt(20)
        obj_font.name = 'Calibri (Body)'
        obj_font.bold = True


        #generate the title page
        my_table = self.my_locating_report.add_table(rows=8, cols=7)
        #merging some cells to format the title page
        a=my_table.cell(0,0)
        b=my_table.cell(7,0)
        ab=a.merge(b)
        
        a=my_table.cell(7,1)
        b=my_table.cell(7,5)
        ab=a.merge(b)

        a=my_table.cell(6,1)
        b=my_table.cell(6,5)
        ab=a.merge(b)

        a=my_table.cell(0,6)
        b=my_table.cell(7,6)
        ab=a.merge(b)


        #vertical line
        for row in range(1,8):
            self.vertical_line(my_table, row, 1)
        #horisontal line
        self.horizontal_line(my_table, 6, 1, width=24)
        

        self.set_cell_margins(my_table, left=10, right=450) 
        


        p = my_table.rows[0].cells[0].add_paragraph()
        r = p.add_run("")
        r.add_picture("Title_image_1.png", height=Mm(195))


        #merging some cells to format the title page
        a=my_table.cell(0,5)
        b=my_table.cell(6,5)
        ab=a.merge(b)  

        p = my_table.rows[0].cells[5].add_paragraph(style="locating_report_style")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run("")
        r.add_picture("Title_image_2.png", width=Mm(80))
        p.add_run(self.templates["Credentials"])
        
        
        my_table.cell(7,1).vertical_alignment = WD_ALIGN_VERTICAL.TOP
        my_table.cell(7,1).text=(self.templates["Title"])
        my_table.rows[7].cells[1].paragraphs[0].style="locating_report_title_style"
        
        
        
        #generate the first page (with vaivers and general info)
               
        current_section = self.my_locating_report.add_section()
        current_section.left_margin = Mm(41)
        current_section.right_margin = Mm(41)
        current_section.top_margin = Mm(45)
        current_section.bottom_margin = Mm(11)
        current_section.header_distance = Mm(10)
        current_section.footer_distance = Mm(10)
        current_section.gutter = Mm(0)

        paragraph_1=self.templates["Intro1"] + self.address + ", " + self.suburb_name+self.templates["Intro2"]
        
        paragraph_2  = self.templates["Conductive_text"]

        paragraph_5  = self.templates["Vaiver"]

        paragraph_10 = self.templates["Operator_text"] + self.operator_name + ", cert. " + self.operator_certificate_number + "."
        
        par1 = self.my_locating_report.add_paragraph(paragraph_1, style="locating_report_style_2")
        par1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        par2= self.my_locating_report.add_paragraph(paragraph_2, style="locating_report_style_2")
        par2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
             
        
        if self.is_GPR.get() !=0:
            paragraph_3 = self.templates["GPR_text"]
            if self.is_GPR_bad_cond.get() !=0:
                paragraph_3+=self.templates["Bad_GPR_conditions"]
            par3= self.my_locating_report.add_paragraph(paragraph_3, style="locating_report_style_2")
            par3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        if self.is_no_non_conductive.get() !=0:
            paragraph_3 = self.templates["No_non_conductive"]
            par3= self.my_locating_report.add_paragraph(paragraph_3, style="locating_report_style_2")
            par3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        if self.is_pushrod.get() !=0:
            paragraph_4 = self.templates["Pushrod_text"]
            par4= self.my_locating_report.add_paragraph(paragraph_4, style="locating_report_style_2")
            par4.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


        par5 = self.my_locating_report.add_paragraph(paragraph_5, style="locating_report_style_2")
        par5.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        par10 = self.my_locating_report.add_paragraph(paragraph_10, style="locating_report_style_2")
        par10.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY



        #generate the pages with a map and photos
        current_section = self.my_locating_report.add_section()
        current_section.left_margin = Mm(13)
        current_section.right_margin = Mm(13)
        current_section.top_margin = Mm(10)
        current_section.bottom_margin = Mm(10)
        current_section.header_distance = Mm(5)
        current_section.footer_distance = Mm(8)
        current_section.gutter = Mm(0)

        #header
        main_header = current_section.header
        main_header.is_linked_to_previous = False
        #generate the table for the header
        my_table = main_header.add_table(rows=3, cols=5, width=Mm(250))
        self.set_cell_margins(my_table, left=0, right=100)
        
        #horisontal line
        for column in range (1,5):
            self.horizontal_line(my_table,0,column)
            self.horizontal_line(my_table,1,column)
        

        #merging some cells to format the header
        a=my_table.cell(0,0)
        b=my_table.cell(2,0)
        ab=a.merge(b)
        a=my_table.cell(0,1)
        b=my_table.cell(0,3)
        ab=a.merge(b)
        a=my_table.cell(1,1)
        b=my_table.cell(1,3)
        ab=a.merge(b)
        a=my_table.cell(2,1)
        b=my_table.cell(2,4)
        ab=a.merge(b)
        
        p = my_table.rows[0].cells[0].add_paragraph(style="locating_report_style")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run()
        r.add_picture("Title_image_2.png", width=Mm(55))


        my_table.style="locating_report_header_style"
        my_table.rows[0].cells[1].text = "Customer: " + self.client_name
        my_table.rows[1].cells[1].text = "Site address: " + self.address + ", " + self.suburb_name
        my_table.rows[0].cells[4].text= "Date: " + self.date
        my_table.rows[1].cells[4].text="Docket#: " + self.docket_number
        my_table.rows[2].cells[1].text = self.templates["General_Notes"]
        my_table.rows[2].cells[4].paragraphs[0].style="locating_report_header_style_2"

        #footer
        main_footer = current_section.footer
        main_footer.is_linked_to_previous = False
       
        p = main_footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style="locating_report_header_style_2"
        r = p.add_run()
        r.add_picture("footer_legend.png", width=Mm(220))

        p2=main_footer.add_paragraph(style="locating_report_header_style_2", text=self.templates["Footer"])
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        

        #main table for the map
        my_table = self.make_main_table()
       
        #map page
        p = my_table.rows[0].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run("")
        r.add_picture("North_arrow.png", width=Mm(17))
        
        
        p = my_table.rows[0].cells[1].paragraphs[0]
        p.style = "locating_report_style_2"
        p.text="Map of the service locating \n"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        #the nested table for the map legend
        my_table = self.make_little_table(my_table.rows[0].cells[1], 6)

        
        p=my_table.rows[0].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.style = "locating_report_style_2"
        r=p.add_run()
        r.add_picture("locating_area.png", width=Mm(8))
        

        p=my_table.rows[0].cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.style = "locating_report_style_2"
        p.add_run(self.templates["Map_legend_1"])


        p=my_table.rows[1].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.style = "locating_report_style_2"
        r=p.add_run()
        r.add_picture("photo_in_report.png", width=Mm(8))
        

        p=my_table.rows[1].cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.style = "locating_report_style_2"
        p.add_run(self.templates["Map_legend_2"])


        #merging some cells to format the bottom of the legend for the map
        a=my_table.cell(5,0)
        b=my_table.cell(5,1)
        a.merge(b)
        
        p=my_table.rows[5].cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.style = "locating_report_style_2"
        r = p.add_run(self.templates["Map_legend_3"])
        

        #split between the tables
        new_par = self.my_locating_report.add_paragraph()
        new_par.style = "locating_report_header_style_2"
        new_par.add_run(" ")



        #making pages for images
        for image_number, image_path in enumerate(self.images_paths):
            my_table = self.make_main_table()
            p = my_table.rows[0].cells[1].paragraphs[0]
            p.style = "locating_report_style_2"
            p.text=f"Photo {image_number+1}" #to start numerating from 1 instead of 0
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            p = my_table.rows[0].cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("\n")
            r = p.add_run()

            #get the dimentions of the image to fit properly:
            img = imread(image_path)
            # Height of the image
            height = img.shape[0]
            # Width of the image
            width = img.shape[1]
            ratio=height/width
            if ratio<0.6:
                r.add_picture(image_path, width=Mm(175))
            else:    
                r.add_picture(image_path, height=Mm(105))
        
            
            #the nested table for the photos legend
            my_table = self.make_little_table(my_table.rows[0].cells[1], 5)
            self.set_cell_margins(my_table) #margins=0
            
            #adding coloured lines
            for line in range(5):
                p=my_table.rows[line].cells[0].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.style = "legend_style_2"
                line_run=p.add_run("___")
                line_run.font.color.rgb = RGBColor(line*40, 254-line*40, line*40) #various colors for each line


            p=my_table.rows[0].cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.style = "legend_style"
            p.add_run(self.templates["Photo_legend_1"])

            p=my_table.rows[1].cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.style = "legend_style"
            p.add_run(self.templates["Photo_legend_2"])


            p=my_table.rows[2].cells[1].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.style = "legend_style"
            p.add_run(self.templates["Photo_legend_3"])


            #split between the tables
            new_par = self.my_locating_report.add_paragraph()
            new_par.style = "locating_report_header_style_2"
            new_par.add_run(" ")

        self.report_file_name = f'Service locating report for {self.address} {self.suburb_name}.docx'

        self.email_template_generator()

        try:
            self.my_locating_report.save(self.report_file_name)
            self.pop_up_message("The report has been created",f"The report and email template have been created successfully. \n The name of the file is \n {self.report_file_name}")
        except PermissionError:
            self.file_save_error_handler(self.my_locating_report,self.report_file_name) 

        
           
        
    def email_template_generator(self):

        my_email_template = Document()
        paragraph_1 = self.templates["email_template_1"] + self.address + ", "+self.suburb_name + f" ({self.date})"
        paragraph_2 = self.templates["email_template_2"] + self.address + ", "+self.suburb_name + f", conducted on {self.date}. " + self.templates["email_template_3"]
        paragraph_2+=self.operator_name + "."
        my_email_template.add_paragraph(paragraph_1)
        my_email_template.add_paragraph(paragraph_2)
        email_template_file_name="Email template for "+self.address + " "+self.suburb_name + ".docx"
        try:
            my_email_template.save(email_template_file_name)
            #self.pop_up_message(f"The file {email_template_file_name} has been saved",f"The file \n {email_template_file_name} \n has been saved successfully.")
        except PermissionError:
            self.file_save_error_handler(my_email_template,email_template_file_name)


    def make_main_table (self):
        

        my_table = self.my_locating_report.add_table(rows=1, cols=2)
        my_table.style = 'Table Grid'
        
        #self.set_cell_margins(my_table, left=Mm(10), right=Mm(10)) #this doesnt work properly here
        
        my_table.autofit = False
        for cell in my_table.columns[0].cells:
            cell.width=Mm(190)
        for cell in my_table.columns[1].cells:
            cell.width=Mm(70)

        my_row = my_table.rows[0]
        my_row.height = Mm(114)
        return my_table
        
        
    def make_little_table(self, root, rows):
        my_table = root.add_table(rows=rows, cols=2)
        #my_table.style = 'Table Grid'
                        
        my_table.autofit = False
        for cell in my_table.columns[0].cells:
            cell.width=Mm(13)
        for cell in my_table.columns[1].cells:
            cell.width=Mm(52)

        my_row = my_table.rows[0]
        my_row.height = Mm(10)
        return my_table
        
                
    def horizontal_line(self, table, row, column, width=7): #bottom line for a cell
        ab=table.cell(row,column)
        tc_pr=ab._element.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
        border_element = OxmlElement('w:bottom')
        border_element.set(qn('w:val'), 'single')
        border_element.set(qn('w:sz'), str(width))
        border_element.set(qn('w:space'), '0')
        border_element.set(qn('w:color'), '000000')
        borders.append(border_element)


    def vertical_line(self, table, row, column, width=24): #right line for a cell
        ab=table.cell(row,column)
        tc_pr=ab._element.get_or_add_tcPr()
        borders = OxmlElement('w:tcBorders')
        tc_pr.append(borders)
        border_element = OxmlElement('w:right')
        border_element.set(qn('w:val'), 'single')
        border_element.set(qn('w:sz'), str(width))
        border_element.set(qn('w:space'), '0')
        border_element.set(qn('w:color'), '000000')
        borders.append(border_element)


    def pop_up_message(self, header_text, main_text):
        popup = Toplevel(self.main_window)
        popup.wm_title(header_text)
        popup.title(main_text)
        popup.geometry("650x200")
        popup.tkraise(self.main_window) # This just tells the message to be on top of the root window.
        Label(popup, text=main_text).pack(side="top", fill="x", pady=10)
        Button(popup, text="Done", command = popup.destroy).pack()
        # Notice that you do not use mainloop() here on the Toplevel() window


    def file_save_error_handler(self, file_content,file_name):
        answer=messagebox.askretrycancel(
            parent=self.main_window,
            title="File saving error - no access",
            message="File saving error. Close the MS Word file and try again"
            )
        
        if answer:
            try:
                file_content.save(file_name)
                self.pop_up_message(f"The file {file_name} has been saved",f"The file \n {file_name} \n has been saved successfully.")
            except PermissionError:
                self.file_save_error_handler(file_content, file_name) 
                
#End of the Class        




my_report=Locating_report()
my_report.load_templates()
my_report.start_window()

